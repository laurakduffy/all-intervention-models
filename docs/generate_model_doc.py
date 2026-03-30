"""Generate a Word (.docx) documentation of the GCR model structure and fund profiles.

Pulls parameter values directly from fund_profiles.py so the documentation
stays in sync when model assumptions change.
"""

import math
import os
import sys

# Ensure fund_profiles.py is importable from gcr-models/
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "gcr-models"))

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from fund_profiles import (
    FUND_PROFILES,
    _AI_CAUSE_FRACTION, _NUCLEAR_CAUSE_FRACTION, _BIO_CAUSE_FRACTION,
    _SENTINEL_REL_REDUCTION_PER_10M, _NUCLEAR_REL_REDUCTION_PER_10M, _AI_REL_REDUCTION_PER_10M,
    _SENTINEL_REL_RISK_REDUCTION, _NUCLEAR_REL_RISK_REDUCTION, _AI_REL_RISK_REDUCTION,
    _TOTAL_XRISK_100YR, _RP_WORLD_PRIORS,
    _r_max_from_cumulative_risk,
)

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "gcr_model_documentation.docx")
M = 10 ** 6

# ---------------------------------------------------------------------------
# Value formatting helpers
# ---------------------------------------------------------------------------

def fmt_dollar(v):
    """Format a dollar amount as $XM or $X."""
    s = f"{v / M:g}"
    return f"${s}M"


def fmt_pct(v, decimals=2):
    """Format a fraction as a percentage string."""
    return f"{v * 100:.{decimals}f}%"


def fmt_sci(v):
    """Format a number in scientific notation, e.g. 1e-7."""
    if v == 0:
        return "0"
    exp = int(math.floor(math.log10(abs(v))))
    mantissa = v / 10 ** exp
    if abs(mantissa - 1.0) < 1e-9:
        return f"1e{exp}"
    return f"{mantissa:.3g}e{exp}"


def fmt_list(vals):
    """Format a plain list as a comma-separated string."""
    return ", ".join(str(v) for v in vals)


def fmt_weighted(entry):
    """Format a weighted sweep entry (dict or plain list) as 'val (X%), ...'."""
    if isinstance(entry, dict):
        vals = entry["values"]
        probs = entry.get("p", [1 / len(vals)] * len(vals))
        return ", ".join(f"{v} ({int(round(p * 100))}%)" for v, p in zip(vals, probs))
    return ", ".join(str(v) for v in entry)


def fmt_r_inf_list(vals):
    """Format r_inf-style list using scientific notation."""
    return ", ".join(fmt_sci(v) for v in vals)


# ---------------------------------------------------------------------------
# Word document helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)


def add_body(doc, text, monospace=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if monospace:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Note: " + text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x50, 0x3C, 0x00)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFADC")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], "D2DCF0")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        bg = "F5F7FC" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cells[i].text = str(val)
            set_cell_bg(cells[i], bg)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Derived values used across sections
# ---------------------------------------------------------------------------

# r_max at central Gaussian params, used in calibration CSV section
_CUM_RISK_SCENARIOS = list(zip(
    ["Low", "Central", "High"],
    ["5%", "10%", "65%"],
    _TOTAL_XRISK_100YR,
    ["Relatively optimistic view of this century",
     "RP house-view baseline",
     "Pessimistic / high-risk scenario"],
))

# r_max values at central Gaussian params (T_peak=15, sigma=33, r_inf=1e-7)
_R_MAX_CENTRAL = {c: _r_max_from_cumulative_risk(c) for _, _, c, _ in _CUM_RISK_SCENARIOS}

_SCENARIO_LABELS = ["conservative", "central", "optimistic"]

_FUND_REL_DATA = {
    "sentinel": {
        "rel_per_10m": _SENTINEL_REL_REDUCTION_PER_10M,
        "rel_rr": _SENTINEL_REL_RISK_REDUCTION,
        "cause_fraction": _BIO_CAUSE_FRACTION,
    },
    "longview_nuclear": {
        "rel_per_10m": _NUCLEAR_REL_REDUCTION_PER_10M,
        "rel_rr": _NUCLEAR_REL_RISK_REDUCTION,
        "cause_fraction": _NUCLEAR_CAUSE_FRACTION,
    },
    "longview_ai": {
        "rel_per_10m": _AI_REL_REDUCTION_PER_10M,
        "rel_rr": _AI_REL_RISK_REDUCTION,
        "cause_fraction": _AI_CAUSE_FRACTION,
    },
}


def _rel_rr_note(fund_key):
    """Build the 'rel_rr_from_int ~ X% of total r_max' note text for a fund."""
    rd = _FUND_REL_DATA[fund_key]
    cause_frac = rd["cause_fraction"]
    rr_ints = [r * cause_frac for r in rd["rel_rr"]["values"]]
    lo = min(rr_ints)
    hi = max(rr_ints)
    return (
        f"rel_rr_from_int = rel_rr * cause_fraction ({cause_frac:.5f}) "
        f"~ {fmt_pct(lo, 4)} to {fmt_pct(hi, 4)} of total r_max."
    )


def _build_rel_rr_rows(fund_key):
    """Build rows for the rel_risk_reduction sweep table of a fund."""
    rd = _FUND_REL_DATA[fund_key]
    rows = []
    for label, rel_per, rel_rr, prob in zip(
        _SCENARIO_LABELS,
        rd["rel_per_10m"]["values"],
        rd["rel_rr"]["values"],
        rd["rel_per_10m"]["p"],
    ):
        rows.append([
            f"{label} ({int(prob * 100)}%)",
            f"{fmt_pct(rel_per, 3)} ({rel_per:.1e})",
            fmt_pct(rel_rr, 4),
        ])
    return rows


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GCR Intervention Model")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fund Profiles: Structure, Assumptions & Model Usage")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x32, 0x5A, 0xA0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Generated from fund_profiles.py and gcr_model.py")
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x78, 0x78, 0x78)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 1. Overview
# ---------------------------------------------------------------------------
add_heading(doc, "1.  Overview", 1)
add_body(doc,
    "This document describes the structure of fund_profiles.py, which contains all "
    "fund-specific assumptions used in the GCR (Global Catastrophic Risk) intervention "
    "valuation model. The model was built by Rethink Priorities to evaluate the cost-"
    "effectiveness of donations to GCR-focused organisations on behalf of Anthropic staff "
    "who are deciding where to direct their personal philanthropy. Rethink Priorities is "
    "the evaluator, not the funder. The model implements a Tarsney-style expected-value "
    "framework that accounts for both near-term and long-term (including stellar expansion) value."
)
add_body(doc, "Three funds are currently modelled:")
add_bullet(doc, f"Sentinel Bio -- biosecurity, extinction-risk pathway + sub-extinction tiers")
add_bullet(doc, f"Longview Nuclear -- nuclear weapons policy, extinction-risk pathway only")
add_bullet(doc, f"Longview AI -- AI safety, extinction-risk pathway only")
add_body(doc,
    "fund_profiles.py is divided into four parts: (1) shared world-prior sweep parameters, "
    "(2) fund-specific parameters, (3) utility functions for loading profiles, and "
    "(4) a __main__ calibration block that prints scenario tables and writes CSVs."
)

# ---------------------------------------------------------------------------
# 2. Model Architecture
# ---------------------------------------------------------------------------
add_heading(doc, "2.  Model Architecture", 1)

add_heading(doc, "2.1  Risk trajectory: the Time-of-Perils Gaussian", 2)
add_body(doc,
    "The model represents annual catastrophic risk as a Gaussian 'time-of-perils' peak "
    "plus a permanent background rate:"
)
add_body(doc, "    r(t)  =  r_inf  +  r_max * exp( -(t - T_peak)^2 / (2*sigma^2) )", monospace=True)
add_body(doc,
    "Parameters T_peak (year_max_risk), sigma (= year_risk_1pct_max / 3), and r_inf are "
    "all swept across scenarios. r_max is derived from the user-specified "
    "cumulative_risk_100_yrs by numerical bisection (see Section 3.5)."
)

add_heading(doc, "2.2  Survival and intervention effect", 2)
add_body(doc,
    "Cumulative survival over T years is the product of (1 - r(t)) for t = 0 ... T. "
    "An intervention reduces the Gaussian component by a fractional amount rel_rr_from_int "
    "during the window [year_effect_starts, year_effect_starts + persistence_effect]:"
)
add_body(doc, "    r_int(t)  =  r_inf  +  r_max * (1 - rel_rr_from_int) * Gaussian(t)", monospace=True)
add_body(doc,
    "The expected value of the intervention is the integral of "
    "(survival_with_int - survival_without_int) * world_value(t) over all future time."
)

add_heading(doc, "2.3  Long-term value", 2)
add_body(doc,
    "World value grows logistically on Earth, then optionally expands cubically via "
    "stellar settlement at time T_c with speed s. The carrying_capacity_multiplier "
    "scales the Earth carrying capacity. T_h = 10^14 years marks the end of the "
    "relevant universe."
)

add_heading(doc, "2.4  Monte Carlo execution", 2)
add_body(doc,
    "run_monte_carlo() in gcr_model.py draws n_samples from the joint distribution "
    "of sweep parameters, runs GCRModel once per sample, then reports percentiles of "
    "total intervention EV. Stratification over cubic_growth, T_c, and r_inf ensures "
    "low-probability strata are sampled proportionally to their assigned weights."
)

# ---------------------------------------------------------------------------
# 3. Shared world priors
# ---------------------------------------------------------------------------
add_heading(doc, "3.  Shared World Prior Parameters (_RP_WORLD_PRIORS)", 1)
add_body(doc,
    "All three funds share the same sweep distribution over world-state parameters. "
    "These are sampled independently of fund-specific parameters in each Monte Carlo draw."
)

add_heading(doc, "3.1  Total x-risk: cumulative_risk_100_yrs", 2)
add_table(doc,
    ["Scenario", "Value", "Meaning"],
    [
        [label, pct_str, meaning]
        for label, pct_str, _, meaning in _CUM_RISK_SCENARIOS
    ],
)
add_body(doc,
    "This is the total x-risk across all causes (AI, bio, nuclear, etc.) over 100 years. "
    "The Gaussian is calibrated to this total so that every simulation includes all causes."
)

add_heading(doc, "3.2  Risk trajectory shape", 2)
_ymr  = _RP_WORLD_PRIORS["year_max_risk"]
_yr1  = _RP_WORLD_PRIORS["year_risk_1pct_max"]
_rinf = _RP_WORLD_PRIORS["r_inf"]
add_table(doc,
    ["Parameter", "Values swept", "Meaning"],
    [
        ["year_max_risk",      fmt_list(_ymr),             "Year of peak annual risk (T_peak)"],
        ["year_risk_1pct_max", fmt_list(_yr1),             "Half-width: year when risk = 1% of peak; sigma = this/3"],
        ["r_inf",              fmt_r_inf_list(_rinf),      "Permanent background annual risk after the peak"],
    ],
)
add_note(doc,
    f"year_risk_1pct_max = {max(_yr1)} means sigma ~ {max(_yr1)//3} years -- a very broad Gaussian. "
    "With high cumulative risk (65%) and a broad Gaussian, the model's exact numerical "
    "r_max solve (Section 3.5) is essential; the old 1.38 approximation would imply "
    ">100% cumulative risk in these scenarios."
)

add_heading(doc, "3.3  Future value trajectory", 2)
_cc   = _RP_WORLD_PRIORS["carrying_capacity_multiplier"]
_rg   = _RP_WORLD_PRIORS["rate_growth"]
_cg   = _RP_WORLD_PRIORS["cubic_growth"]
_tc   = _RP_WORLD_PRIORS["T_c"]
_s    = _RP_WORLD_PRIORS["s"]
add_table(doc,
    ["Parameter", "Values (weights)", "Meaning"],
    [
        ["carrying_capacity_multiplier", fmt_weighted(_cc), "Earth carrying capacity as multiple of initial_value"],
        ["rate_growth",                  fmt_list(_rg),     "Logistic growth rate of world value"],
        ["cubic_growth",                 fmt_weighted(_cg), "Whether stellar settlement occurs"],
        ["T_c",                          fmt_weighted(_tc), "Year stellar settlement begins (if cubic_growth=True)"],
        ["s",                            fmt_list(_s),      "Speed of stellar expansion"],
    ],
)
_tc_modal = _tc["values"][_tc["p"].index(max(_tc["p"]))]
_tc_modal_pct = int(round(max(_tc["p"]) * 100))
add_body(doc,
    f"Sampling weights are non-uniform for cubic_growth and T_c: stellar settlement "
    f"within {min(_tc['values'])} years is considered unlikely "
    f"({int(round(_tc['p'][_tc['values'].index(min(_tc['values']))] * 100))}%), "
    f"and T_c = {_tc_modal} is the modal view ({_tc_modal_pct}%). "
    "Stratification in run_monte_carlo() allocates MC samples in proportion to the "
    "product of marginal probabilities, ensuring rare strata are still represented."
)

add_heading(doc, "3.4  Fixed value parameters (all funds)", 2)
# Use sentinel as a representative fund for shared fixed params
_fp_ref = FUND_PROFILES["sentinel"]["fixed_params"]
add_table(doc,
    ["Parameter", "Value", "Meaning"],
    [
        ["initial_value", f"{_fp_ref['initial_value']:.2e}", "Current world value (~present-day proxy, QALYs or welfare units)"],
        ["T_h",           f"10^14 years",                    "Time horizon: end of relevant universe"],
        ["periods_value", str(_fp_ref["periods_value"]),     "Reporting breakpoints for short-term EV decomposition"],
    ],
)

add_heading(doc, "3.5  r_max: numerical solve", 2)
add_body(doc,
    "r_max (the Gaussian peak annual risk) is not a free parameter -- it is derived. "
    "Given cumulative_risk_100_yrs, year_max_risk, year_risk_1pct_max, and r_inf, the model "
    "solves numerically for the unique r_max satisfying:"
)
add_body(doc,
    "    1 - prod_{t=0}^{100} (1 - clip(r_inf + r_max * G_t, 0, 1))  =  cumulative_risk_100_yrs",
    monospace=True,
)
add_body(doc,
    "This uses 60 iterations of vectorized bisection (~10^-18 precision), one solve "
    "per Monte Carlo sample. The old closed-form 1.38 * annual_rate approximation is "
    "no longer used in the model (only retained in the calibration CSV as a rough check)."
)

# ---------------------------------------------------------------------------
# 4. Fund-specific parameters
# ---------------------------------------------------------------------------
add_heading(doc, "4.  Fund-Specific Parameters", 1)

add_heading(doc, "4.1  The rel_risk_reduction parameterisation (Option A)", 2)
add_body(doc,
    "Each fund specifies how much of the total Gaussian peak risk it reduces, as a "
    "fraction, per dollar of budget. This is the key fund-effectiveness parameter. "
    "The model computes:"
)
add_body(doc, "    rel_rr_from_int  =  rel_risk_reduction * cause_fraction", monospace=True)
add_body(doc, "    r_int(t)         =  r_inf + r_max * (1 - rel_rr_from_int) * Gaussian(t)", monospace=True)
add_body(doc,
    "rel_risk_reduction is swept as a fraction of cause-specific r_max per unit budget, "
    "independent of the cumulative_risk_100_yrs scenario. This avoids a Cartesian-product "
    "inconsistency that arose under the old abs_risk_reduction parameterisation, where the "
    "intervention magnitude was correlated with the risk level being swept."
)

add_heading(doc, "4.2  cause_fraction: attributing total risk to each cause", 2)
add_table(doc,
    ["Fund", "Formula", "Value", "Interpretation"],
    [
        ["Sentinel Bio",     "(0.004183) / 0.67",         f"{_BIO_CAUSE_FRACTION:.5f}",  "Bio share of total x-risk (RP Cross Cause Model)"],
        ["Longview Nuclear", "(0.02354) / 0.67",          f"{_NUCLEAR_CAUSE_FRACTION:.5f}", "Nuclear share of total x-risk"],
        ["Longview AI",      "(0.5541 + 0.06157) / 0.67", f"{_AI_CAUSE_FRACTION:.5f}",   "AI direct + AI indirect share"],
    ],
)
add_body(doc,
    "The denominator 0.67 is the total probability mass assigned to modelled causes "
    "in the RP Cross Cause Model. cause_fraction * rel_risk_reduction gives the fraction "
    "of the total Gaussian peak (r_max) reduced by the intervention."
)

# ---------------------------------------------------------------------------
# 5. Per-fund profiles
# ---------------------------------------------------------------------------
add_heading(doc, "5.  Per-Fund Profiles", 1)


def _add_fund_section(doc, fund_key):
    """Add a full per-fund section, pulling all values from FUND_PROFILES."""
    profile  = FUND_PROFILES[fund_key]
    fp       = profile["fixed_params"]
    budget   = profile["budget"]

    add_heading(doc, "Budget and counterfactual", 3)
    add_table(doc,
        ["Parameter", "Value", "Source / reasoning"],
        [
            ["budget",                fmt_dollar(budget),                        "Grant amount being evaluated"],
            ["counterfactual_factor", f"{profile['counterfactual_factor']:.3f}", "Weighted counterfactual scenarios"],
            ["p_harm",                fmt_pct(profile["p_harm"], 0),             "Probability intervention is counterproductive"],
            ["p_zero",                fmt_pct(profile["p_zero"], 0),             "Probability intervention has no extinction-risk effect"],
            ["harm_multiplier",       str(profile["harm_multiplier"]),           "Harm effect magnitude relative to benefit"],
        ],
    )

    add_heading(doc, "Intervention timing", 3)
    add_table(doc,
        ["Parameter", "Value", "Source / reasoning"],
        [
            ["year_effect_starts", str(fp["year_effect_starts"]), "Year intervention begins taking effect"],
            ["persistence_effect", str(fp["persistence_effect"]), "Years the risk reduction persists"],
        ],
    )

    add_heading(doc, "Relative risk reduction sweep", 3)
    budget_label = fmt_dollar(budget)
    rows = _build_rel_rr_rows(fund_key)
    add_table(doc,
        ["Scenario", "rel per $10M", f"rel_rr at {budget_label} budget", "Source"],
        [row + ["Field-level reasoning"] for row in rows],
    )
    add_note(doc, _rel_rr_note(fund_key))


# --- Sentinel Bio ---
add_heading(doc, "5.1  Sentinel Bio", 2)
_add_fund_section(doc, "sentinel")

add_heading(doc, "Sub-extinction tiers", 3)
add_body(doc,
    "All three funds include sub-extinction modelling (recoverable catastrophes). "
    "Two tiers use the simple EV formula: P(event/yr) * deaths * rel_rr * persistence * counterfactual."
)
for _sub_fund_key in ("sentinel", "longview_nuclear", "longview_ai"):
    _tiers = FUND_PROFILES[_sub_fund_key].get("sub_extinction_tiers", [])
    if not _tiers:
        continue
    add_heading(doc, FUND_PROFILES[_sub_fund_key]["display_name"], 4)
    add_table(doc,
        ["Tier", "P(event/10yr)", "Expected deaths", "discount"],
        [
            [
                t["tier_name"],
                fmt_pct(t["p_10yr"], 0),
                f"{t['expected_deaths'] / 1e6:.1f}M (geomean of tier bounds)",
                f"{t['discount']:.2g} -- "
                + ("no discount" if t["discount"] == 1.0 else f"{fmt_pct(1 - t['discount'], 0)} discount"),
            ]
            for t in _tiers
        ],
    )

# --- Longview Nuclear ---
add_heading(doc, "5.2  Longview Nuclear", 2)
_add_fund_section(doc, "longview_nuclear")
add_note(doc,
    f"The 10x discount reflects that Q4.4 responses (0.2% per $10M) were judged optimistic. "
    f"nuclear cause_fraction = {_NUCLEAR_CAUSE_FRACTION:.5f}, so "
    f"rel_rr_from_int ranges from "
    f"{fmt_pct(min(_NUCLEAR_REL_RISK_REDUCTION['values']) * _NUCLEAR_CAUSE_FRACTION, 5)} to "
    f"{fmt_pct(max(_NUCLEAR_REL_RISK_REDUCTION['values']) * _NUCLEAR_CAUSE_FRACTION, 5)} of total r_max."
)

# --- Longview AI ---
add_heading(doc, "5.3  Longview AI", 2)
_add_fund_section(doc, "longview_ai")
_ai_rr_ints = [r * _AI_CAUSE_FRACTION for r in _AI_REL_RISK_REDUCTION["values"]]
_r_max_10pct = _R_MAX_CENTRAL[0.10]
add_note(doc,
    f"AI cause_fraction = {_AI_CAUSE_FRACTION:.3f}, so rel_rr_from_int = "
    f"{fmt_pct(min(_ai_rr_ints), 3)} to {fmt_pct(max(_ai_rr_ints), 3)} of total r_max. "
    f"At 10% cumulative x-risk (r_max ~ {_r_max_10pct:.5f}) this corresponds to roughly "
    f"{min(_ai_rr_ints) * _r_max_10pct * 1e4:.0f}–"
    f"{max(_ai_rr_ints) * _r_max_10pct * 1e4:.0f} bp absolute risk reduction "
    f"per $1B (peak-annual basis)."
)

# ---------------------------------------------------------------------------
# 6. Parameter flow
# ---------------------------------------------------------------------------
add_heading(doc, "6.  How Parameters Flow into the Model", 1)

add_heading(doc, "6.1  Step-by-step flow", 2)
steps = [
    ("fund_profiles.py",   "Defines sweep_params (sampled per MC draw) and fixed_params (constant)."),
    ("run_monte_carlo()",  "Draws n_samples from sweep_params; stratifies by cubic_growth, T_c, r_inf; applies p_harm/p_zero adjustments."),
    ("GCRParams",          "Receives vectorised arrays (one element per sample) for all parameters."),
    ("GCRModel._derive()", "Calls _solve_r_max() to find r_max per sample; computes rel_rr_from_int = rel_risk_reduction * cause_fraction."),
    ("GCRModel.run()",     "Computes annual risk r(t) with and without intervention, survival products, and value-weighted EV integrals by period."),
    ("Output",             "Returns total_values[n_samples] and percentiles (p1...p99, mean)."),
]
for label, desc in steps:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{label}: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

doc.add_paragraph()
add_heading(doc, "6.2  Key formula chain", 2)
add_body(doc, "Starting from fund inputs to final intervention EV per dollar:")
formulas = [
    "rel_risk_reduction       [swept, e.g. 0.0007]",
    "  * cause_fraction       [fixed per fund, e.g. {:.3f} for AI]".format(_AI_CAUSE_FRACTION),
    "  = rel_rr_from_int      [fraction of total r_max reduced]",
    "  * r_max                [peak annual risk, solved from cumulative_risk_100_yrs]",
    "  = delta_r(t)           [absolute annual risk reduction at peak]",
    "  * Gaussian(t)          [applied only during intervention window]",
    "  -> delta_survival(t)   [difference in survival probability]",
    "  * world_value(t)       [value of surviving year t]",
    "  = EV contribution      [integrated over all t, summed = total EV]",
    "  / budget               [normalised to EV per dollar]",
]
for f in formulas:
    add_body(doc, f, monospace=True)

# ---------------------------------------------------------------------------
# 7. Calibration CSV
# ---------------------------------------------------------------------------
add_heading(doc, "7.  Calibration CSV Output", 1)
add_body(doc,
    "Running fund_profiles.py as __main__ writes two CSV files to the gcr-models directory:"
)
add_bullet(doc,
    "calibration_abs_risk_reduction_detail.csv -- one row per "
    "(fund x rel_scenario x cum_risk_scenario) combination (27 rows total)."
)
add_bullet(doc,
    "calibration_abs_risk_reduction_summary.csv -- one row per fund with "
    "min / max / mean / median / geometric_mean of peak_annual_abs_risk_reduction_bp_per_1b."
)
doc.add_paragraph()
add_body(doc,
    "The key metric is peak_annual_abs_risk_reduction_bp_per_1b: the absolute reduction in "
    "peak annual risk per $1B spent, in basis points (1 bp = 0.01%). Formula:"
)
add_body(doc,
    "    peak_annual_abs_rr_bp_per_1b  =  rel_per_10m * cause_fraction * r_max * 1,000,000",
    monospace=True,
)
add_body(doc,
    "The factor 1,000,000 = 100 (scaling $10M to $1B) * 10,000 (fraction to basis points). "
    "This uses peak annual risk (r_max) rather than cumulative risk, because the "
    "intervention's persistence is far shorter than 100 years."
)

# r_max note: use central Gaussian params (T_peak=15, sigma=33, r_inf=1e-7)
_r_max_note_parts = [
    f"r_max({pct_str}) ~ {r:.6f}"
    for _, pct_str, cum, _ in _CUM_RISK_SCENARIOS
    for r in [_R_MAX_CENTRAL[cum]]
]
add_note(doc,
    "r_max values used in CSV are solved for central Gaussian params "
    f"(T_peak={_RP_WORLD_PRIORS['year_max_risk'][1]}, "
    f"sigma={_RP_WORLD_PRIORS['year_risk_1pct_max'][1]//3}, "
    f"r_inf=1e-7): "
    + ", ".join(_r_max_note_parts) + "."
)

# Compute summary stats from the same logic as fund_profiles.py __main__
_CSV_FUND_CONFIGS = [
    {"fund": "Sentinel Bio",    "rel_per_10m": _SENTINEL_REL_REDUCTION_PER_10M, "cause_fraction": _BIO_CAUSE_FRACTION},
    {"fund": "Longview Nuclear","rel_per_10m": _NUCLEAR_REL_REDUCTION_PER_10M,  "cause_fraction": _NUCLEAR_CAUSE_FRACTION},
    {"fund": "Longview AI",     "rel_per_10m": _AI_REL_REDUCTION_PER_10M,       "cause_fraction": _AI_CAUSE_FRACTION},
]

def _geo_mean(vals):
    return math.exp(sum(math.log(v) for v in vals) / len(vals))

calib_rows = []
for cfg in _CSV_FUND_CONFIGS:
    bp_vals = [
        rel * cfg["cause_fraction"] * r_max * 1_000_000
        for rel in cfg["rel_per_10m"]["values"]
        for r_max in _R_MAX_CENTRAL.values()
    ]
    calib_rows.append([
        cfg["fund"],
        f"{min(bp_vals):.3f}",
        f"{max(bp_vals):.3f}",
        f"{_geo_mean(bp_vals):.3f}",
    ])

add_table(doc,
    ["Fund", "min bp/bn", "max bp/bn", "geo-mean bp/bn"],
    calib_rows,
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT_PATH)
print(f"Written: {OUT_PATH}")
